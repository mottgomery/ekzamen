from flask import Flask, render_template, redirect, url_for, flash, abort, send_file, send_from_directory, session
from flask_sqlalchemy import SQLAlchemy
from flask_bcrypt import Bcrypt
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
import os
from base64 import b64encode
from docx import Document
import docx.shared
from io import BytesIO
from pytz import timezone


# Функция для конвертации времени
def format_time_to_local(dt, tz_name="Europe/Moscow"):
    local_tz = timezone(tz_name)
    local_time = dt.astimezone(local_tz)
    return local_time.strftime('%d.%m.%Y %H:%M')








app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///zadachi.db'
app.config['SECRET_KEY'] = 'your_secret_key'
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.jinja_env.filters['format_time_to_local'] = format_time_to_local

db = SQLAlchemy(app)
bcrypt = Bcrypt(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'


# Модель тега
class Tag(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(50), unique=True, nullable=False)

# Таблица связывания задач и тегов
post_tag = db.Table('post_tag',
    db.Column('post_id', db.Integer, db.ForeignKey('post.id'), primary_key=True),
    db.Column('tag_id', db.Integer, db.ForeignKey('tag.id'), primary_key=True)
)

# Обновление модели задачи для поддержки тегов
class Post(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(300), nullable=False)
    zad = db.Column(db.LargeBinary, nullable=False)
    resh1 = db.Column(db.LargeBinary, nullable=False)
    resh2 = db.Column(db.LargeBinary, nullable=True)
    resh3 = db.Column(db.LargeBinary, nullable=True)
    resh4 = db.Column(db.LargeBinary, nullable=True)
    resh5 = db.Column(db.LargeBinary, nullable=True)
    level = db.Column(db.Integer, nullable=False)
    tags = db.relationship('Tag', secondary=post_tag, backref=db.backref('posts', lazy='dynamic'))

# Модель пользователя
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(150), unique=True, nullable=False)
    password = db.Column(db.String(150), nullable=False)
    role = db.Column(db.String(50), nullable=False)  # 'admin', 'teacher', 'student'

# Добавим новую модель для тестов
class Test(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(300), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    user = db.relationship('User', backref=db.backref('tests', lazy=True))  # Связь с моделью User
    tasks = db.relationship('Post', secondary='test_post', backref=db.backref('tests', lazy='dynamic'))

# Ассоциативная таблица для задач и тестов
test_post = db.Table('test_post',
    db.Column('test_id', db.Integer, db.ForeignKey('test.id'), primary_key=True),
    db.Column('post_id', db.Integer, db.ForeignKey('post.id'), primary_key=True)
)



class Thread(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(300), nullable=False)
    votes = db.Column(db.Integer, default=0)
    is_admin_thread = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=db.func.current_timestamp())
    comments = db.relationship("ThreadPost", backref="thread", cascade="all, delete-orphan", lazy=True)

class ThreadPost(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    thread_id = db.Column(db.Integer, db.ForeignKey("thread.id"), nullable=False)
    name = db.Column(db.String(100), default="Аноним")
    message = db.Column(db.Text, nullable=False)
    file = db.Column(db.LargeBinary, nullable=True)  # Для прикрепленных файлов
    file_name = db.Column(db.String(300), nullable=True)
    created_at = db.Column(db.DateTime, default=db.func.current_timestamp())
    replies = db.relationship("Reply", backref="parent", cascade="all, delete-orphan", lazy=True)


class Reply(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    post_id = db.Column(db.Integer, db.ForeignKey("thread_post.id"), nullable=False)
    file = db.Column(db.LargeBinary, nullable=True)  # Для прикрепленных файлов
    file_name = db.Column(db.String(300), nullable=True)
    name = db.Column(db.String(100), default="Аноним")
    message = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=db.func.current_timestamp())

@app.route('/forum')
def forum():
    threads = Thread.query.order_by(Thread.votes.desc()).all()
    return render_template('forum.html', threads=threads)


@app.route('/thread/<int:thread_id>', methods=['GET', 'POST'])
def view_thread(thread_id):
    thread = Thread.query.get_or_404(thread_id)
    page = request.args.get('page', 1, type=int)
    comments_paginated = ThreadPost.query.filter_by(thread_id=thread.id).order_by(ThreadPost.created_at.desc()).paginate(page=page, per_page=50, error_out=False)

    if request.method == 'POST':
        if thread.is_admin_thread and current_user.role != 'admin':
            abort(403)

        message = request.form['message']
        name = request.form.get('name', 'Аноним')
        parent_id = request.form.get('parent_id')

        # Получаем файл
        file = request.files.get('file')
        file_data, file_name = (None, None)

        if file and file.filename:
            file_data = file.read()
            file_name = file.filename

        if parent_id:
            reply = Reply(post_id=parent_id, message=message, name=name, file=file_data, file_name=file_name)
            db.session.add(reply)
        else:
            post = ThreadPost(thread_id=thread_id, message=message, name=name, file=file_data, file_name=file_name)
            db.session.add(post)

        db.session.commit()
        return redirect(url_for('view_thread', thread_id=thread_id, page=page))

    return render_template('thread.html', thread=thread, comments_paginated=comments_paginated)



@app.route('/new_admin_thread', methods=['POST'])
@login_required
def new_admin_thread():
    if current_user.role != 'admin':
        abort(403)

    title = request.form['title']
    thread = Thread(title=title, votes=0, is_admin_thread=True)
    db.session.add(thread)
    db.session.commit()
    return redirect(url_for('forum'))


@app.route('/forum/remove_comment/<int:comment_id>', methods=['POST'])
@login_required
def delete_comment(comment_id):
    comment = ThreadPost.query.get_or_404(comment_id)
    if current_user.role != 'admin':
        flash('Только администратор может удалять комментарии.', 'danger')
        return redirect(url_for('view_thread', thread_id=comment.thread_id))
    db.session.delete(comment)
    db.session.commit()
    flash('Комментарий и все ответы удалены.', 'success')
    return redirect(url_for('view_thread', thread_id=comment.thread_id))

@app.route('/new_thread', methods=['POST'])
def new_thread():
    title = request.form['title']
    thread = Thread(title=title, votes=0)
    db.session.add(thread)
    db.session.commit()
    return redirect(url_for('forum'))


import json
from flask import request, jsonify, make_response

@app.route('/accept-cookies')
def accept_cookies():
    response = make_response("Cookies accepted!")
    response.set_cookie('accepted_cookies', 'true', max_age=60*60*24*100)  # 100 дней
    return response





@app.route('/vote_thread/<int:thread_id>/<action>', methods=['POST'])
@login_required  # Требует авторизации
def vote_thread(thread_id, action):
    thread = db.session.get(Thread, thread_id)
    if not thread:
        return jsonify({"error": "Тред не найден."}), 404

    if 'user_votes' not in session:
        session['user_votes'] = {}

    user_votes = session['user_votes']

    if str(thread_id) in user_votes:
        return jsonify({"error": "Вы уже голосовали за этот тред."}), 400

    if action == 'up':
        thread.votes += 1
    elif action == 'down':
        thread.votes -= 1
    else:
        return jsonify({"error": "Неверное действие."}), 400

    user_votes[str(thread_id)] = action
    session['user_votes'] = user_votes
    session.modified = True

    db.session.commit()
    return jsonify({"votes": thread.votes})



@app.route('/vote_post/<int:post_id>/<action>', methods=['POST'])
def vote_post(post_id, action):
    post = Post.query.get_or_404(post_id)
    if action == 'up':
        post.votes += 1
    elif action == 'down':
        post.votes -= 1
    db.session.commit()
    return jsonify({'votes': post.votes})


@app.route('/forum/remove_thread/<int:thread_id>', methods=['POST'])
@login_required
def delete_thread(thread_id):
    thread = Thread.query.get_or_404(thread_id)
    if current_user.role != 'admin':
        flash('Только администратор может удалять треды.', 'danger')
        return redirect(url_for('forum'))
    db.session.delete(thread)
    db.session.commit()
    flash("Тред и все связанные комментарии и ответы удалены.", "success")
    return redirect(url_for("forum"))


@app.route('/download_post_file/<int:post_id>')
def download_post_file(post_id):
    post = ThreadPost.query.get_or_404(post_id)

    if not post.file or not post.file_name:
        abort(404)

    mime_type = "application/octet-stream"
    if post.file_name.endswith(('.png', '.jpg', '.jpeg', '.gif')):
        mime_type = "image/png"  # Можно улучшить определение MIME-типов

    return send_file(
        BytesIO(post.file),
        mimetype=mime_type,
        as_attachment=not post.file_name.endswith(('.png', '.jpg', '.jpeg', '.gif')),
        download_name=post.file_name
    )


@app.route('/search', methods=['GET'])
def search():
    query = request.args.get('q', '').strip()
    if not query:
        return render_template('search.html', query=query, results=[])

    # Поиск по задачам и тестам
    task_results = Post.query.filter(Post.title.ilike(f"%{query}%")).all()
    test_results = Test.query.filter(Test.name.ilike(f"%{query}%")).all()

    # Передача результатов в шаблон
    return render_template('search.html', query=query, task_results=task_results, test_results=test_results)




@app.route('/create_test', methods=['POST'])
@login_required
def create_test():
    data = request.get_json()
    test_name = data.get('testName')
    task_ids = data.get('taskIds')

    if not test_name or not task_ids:
        return jsonify({'error': 'Invalid input'}), 400

    tasks = Post.query.filter(Post.id.in_(task_ids)).all()
    if len(tasks) != len(task_ids):
        return jsonify({'error': 'Some tasks not found'}), 404

    new_test = Test(name=test_name, user_id=current_user.id, tasks=tasks)
    db.session.add(new_test)
    db.session.commit()

    return jsonify({'success': True}), 201

@app.route('/tests', methods=['GET'])
@login_required
def view_tests():
    page = request.args.get('page', 1, type=int)
    tests = Test.query.filter_by(user_id=current_user.id).paginate(page=page, per_page=10)
    posts = Post.query.paginate(page=page, per_page=3000)
    return render_template('mytest.html', tests=tests, posts=posts)

@app.route('/all_tests', methods=['GET'])
def all_tests():
    page = request.args.get('page', 1, type=int)
    tests = Test.query.paginate(page=page, per_page=10)  # Пагинация для всех тестов
    return render_template('all_tests.html', tests=tests)

# Страница отображения конкретного теста
@app.route('/test_details/<int:test_id>')
def show_test_detail(test_id):
    test = Test.query.get_or_404(test_id)
    return render_template('test_detail.html', test=test)


@app.route('/delete_test/<int:test_id>', methods=['POST'])
@login_required
def delete_test(test_id):
    test = Test.query.get_or_404(test_id)
    if test.user_id != current_user.id:
        flash('Вы не можете удалить этот тест.', 'danger')
        return redirect(url_for('view_tests'))

    try:
        db.session.delete(test)
        db.session.commit()
        flash('Тест успешно удалён.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Ошибка при удалении теста: {e}', 'danger')

    return redirect(url_for('view_tests'))





@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username'].strip()
        password = request.form['password']
        user = User.query.filter_by(username=username).first()

        # Отладочные сообщения
        print(f"Username: {username}")
        print(f"User found: {user.username if user else 'None'}")

        if user and bcrypt.check_password_hash(user.password, password):
            print("Password match: True")
            login_user(user)
            return redirect(url_for('index'))
        else:
            print("Password match: False")
            flash('Неправильные имя пользователя или пароль.')

    return render_template('login.html')



@app.route('/logout', methods=['GET', 'POST'])
def logout():
    # Логика выхода пользователя
    logout_user()  # вызов функции из Flask-Login
    return redirect(url_for('index'))  # перенаправление на главную страницу


@app.route('/register_teacher', methods=['GET', 'POST'])
@login_required
def register_teacher():
    if current_user.role != 'admin':
        flash('Доступ запрещен.')
        return redirect(url_for('index'))

    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
        new_user = User(username=username, password=hashed_password, role='teacher')

        try:
            db.session.add(new_user)
            db.session.commit()
            flash('Учитель успешно добавлен.')
            return redirect(url_for('index'))
        except Exception as e:
            flash(f'Ошибка при добавлении учителя: {e}')

    tags = Tag.query.all()
    return render_template('register_teacher.html', tags=tags)







@app.route('/generate_docx/<int:test_id>', methods=['GET'])
@login_required
def generate_docx(test_id):
    test = Test.query.get_or_404(test_id)
    if test.user_id != current_user.id:
        abort(403)

    doc = Document()
    doc.add_heading(test.name, level=1)


    for i, task in enumerate(test.tasks, start=1):
        doc.add_heading(f"№ {i}", level=2)

        try:
            if task.zad:
                image_stream = BytesIO(task.zad)
                doc.add_paragraph("\n", style="Normal")  # Отступ для выравнивания изображения
                run = doc.paragraphs[-1].add_run()
                run.add_picture(image_stream, width=docx.shared.Inches(4))  # Уменьшение размера изображения
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = 1  # Центровка изображения
            else:
                doc.add_paragraph("Нет изображения для этой задачи.")
        except Exception as e:
            doc.add_paragraph(f"Ошибка при добавлении изображения задачи {i}: {str(e)}")

        if hasattr(task, 'description') and task.description:
            doc.add_paragraph(task.description)

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(output, as_attachment=True, download_name=f"{test.name}.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')



@app.route('/add_tag', methods=['POST'])
@login_required
def add_tag():
    if current_user.role != 'admin':
        flash('Только администратор может добавлять теги.')
        return redirect(url_for('index'))

    tag_name = request.form.get('tag_name').strip().lower()
    if not tag_name:
        flash('Введите название тега.')
        return redirect(url_for('register_teacher'))

    existing_tag = Tag.query.filter_by(name=tag_name).first()
    if existing_tag:
        flash('Такой тег уже существует.')
        return redirect(url_for('register_teacher'))

    new_tag = Tag(name=tag_name)
    db.session.add(new_tag)
    db.session.commit()
    flash('Тег успешно добавлен.')
    return redirect(url_for('register_teacher'))


@app.route('/delete_tags', methods=['POST'])
@login_required
def delete_tags():
    if current_user.role != 'admin':
        flash('Только администратор может удалять теги.')
        return redirect(url_for('register_teacher'))

    tags_to_delete = request.form.getlist('tags_to_delete')  # Список выбранных ID тегов
    if not tags_to_delete:
        flash('Выберите хотя бы один тег для удаления.')
        return redirect(url_for('register_teacher'))

    try:
        for tag_id in tags_to_delete:
            tag = Tag.query.get(tag_id)
            if tag:
                db.session.delete(tag)
        db.session.commit()
        flash('Выбранные теги успешно удалены.')
    except Exception as e:
        db.session.rollback()
        flash(f'Ошибка при удалении тегов: {e}')

    return redirect(url_for('register_teacher'))


@app.route('/delete_solution/<int:post_id>/<int:solution_number>', methods=['POST', 'GET'])
@login_required
def delete_solution(post_id, solution_number):
    post = Post.query.get_or_404(post_id)

    # Проверяем права доступа
    if current_user.role not in ['admin', 'teacher']:
        abort(403)

    # Удаляем выбранное решение
    if solution_number == 1:
        post.resh1 = None
    elif solution_number == 2:
        post.resh2 = None
    elif solution_number == 3:
        post.resh3 = None
    elif solution_number == 4:
        post.resh3 = None
    elif solution_number == 5:
        post.resh3 = None
    else:
        abort(400, description="Неверный номер решения")

    db.session.commit()
    flash('Решение удалено', 'success')
    return redirect(url_for('show_post_detail', id=post_id))


@app.route('/post/<int:post_id>/edit_tags', methods=['GET', 'POST'])
def update_post_tags(post_id):  # Уникальное имя функции
    post = Post.query.get_or_404(post_id)
    all_tags = Tag.query.all()  # Получаем все доступные теги

    if request.method == 'POST':
        # Получаем выбранные теги из формы
        selected_tags_ids = request.form.getlist('tags')  # Список ID выбранных тегов
        selected_tags = Tag.query.filter(Tag.id.in_(selected_tags_ids)).all()

        # Обновляем теги задачи
        post.tags = selected_tags
        db.session.commit()
        flash('Теги успешно обновлены.', 'success')
        return redirect(url_for('show_post_detail', id=post.id))


    return render_template('edit_tags.html', post=post, all_tags=all_tags)




@app.route('/create', methods=['GET', 'POST'])
@login_required
def create():
    if current_user.role not in ['admin', 'teacher']:
        flash('Только администраторы и учителя могут добавлять задачи.')
        return redirect(url_for('index'))

    if request.method == 'POST':
        title = request.form['title']
        level = request.form['level']
        tag_ids = request.form.getlist('tags')  # Получаем выбранные теги

        # Загружаем файлы
        zad = request.files['zad'].read()
        resh1 = request.files['resh1'].read()
        resh2 = request.files['resh2'].read() if 'resh2' in request.files else None
        resh3 = request.files['resh3'].read() if 'resh3' in request.files else None
        resh4 = request.files['resh4'].read() if 'resh4' in request.files else None
        resh5 = request.files['resh5'].read() if 'resh5' in request.files else None

        # Создаем задачу
        post = Post(title=title, zad=zad, resh1=resh1, resh2=resh2, resh3=resh3, resh4=resh4, resh5=resh5, level=level)

        # Добавляем выбранные теги к задаче
        for tag_id in tag_ids:
            tag = Tag.query.get(int(tag_id))
            if tag:
                post.tags.append(tag)

        try:
            db.session.add(post)
            db.session.commit()
            flash('Задача успешно добавлена.')
            return redirect('/')
        except Exception as e:
            flash(f'Ошибка при добавлении задачи: {e}')
            db.session.rollback()

    all_tags = Tag.query.all()
    return render_template('create.html', all_tags=all_tags)

@app.route('/update_solutions/<int:post_id>', methods=['POST'])
@login_required
def update_solutions(post_id):
    post = Post.query.get_or_404(post_id)
    if current_user.role not in ['admin', 'teacher']:
        abort(403)

    # Обновление решений
    if 'resh1' in request.files:
        resh1 = request.files['resh1']
        if resh1.filename:
            post.resh1 = resh1.read()

    if 'resh2' in request.files:
        resh2 = request.files['resh2']
        if resh2.filename:
            post.resh2 = resh2.read()

    if 'resh3' in request.files:
        resh3 = request.files['resh3']
        if resh3.filename:
            post.resh3 = resh3.read()

    if 'resh4' in request.files:
        resh4 = request.files['resh3']
        if resh4.filename:
            post.resh4 = resh4.read()

    if 'resh5' in request.files:
        resh5 = request.files['resh3']
        if resh5.filename:
            post.resh5 = resh5.read()


    db.session.commit()
    flash('Решения успешно обновлены!', 'success')
    return redirect(url_for('show_post_detail', id=post_id))


# Главная страница
@app.route('/')
@app.route('/index')
def index():
    # Получаем количество записей в каждой таблице
    tasks_count = Post.query.count()
    threads_count = Thread.query.count()
    users_count = User.query.count()
    tests_count = Test.query.count()
    accepted_cookies = request.cookies.get('accepted_cookies')
    return render_template('index.html', tasks_count=tasks_count, threads_count=threads_count, users_count=users_count, tests_count=tests_count, accepted_cookies=accepted_cookies)

@app.route('/delete/<int:id>', methods=['POST'])
def delete_post(id):
    post = Post.query.get_or_404(id)

    try:
        db.session.delete(post)
        db.session.commit()
        return redirect('/posts')
    except Exception as e:
        return f'Ошибка при удалении задачи: {e}'


@app.route('/post/<int:id>/edit_tags', methods=['POST'])
@login_required
def edit_tags(id):
    post = Post.query.get_or_404(id)

    # Проверяем, что текущий пользователь — учитель или администратор
    if current_user.role not in ['admin', 'teacher']:
        flash('Недостаточно прав для редактирования тегов.')
        return redirect(url_for('show_post_detail', id=id))

    # Обновляем теги задачи
    tag_ids = request.form.getlist('tags')
    post.tags = [Tag.query.get(int(tag_id)) for tag_id in tag_ids if Tag.query.get(int(tag_id))]

    db.session.commit()
    flash('Теги успешно обновлены.')
    return redirect(url_for('show_post_detail', id=id))

@app.route('/clear-cookies')
def clear_cookies():
    response = make_response("Cookies cleared!")
    response.set_cookie('accepted_cookies', '', max_age=0)  # Очистка cookie
    return response


# Страница со всеми задачами
@app.route('/posts', methods=["GET"])
def posts():
    sort_by = request.args.get('sort_by', 'id')
    order = request.args.get('order', 'asc')
    filter_tags = request.args.getlist('tags')
    selected_tag = request.args.get('tags')
    filter_title = request.args.get('title')

    query = Post.query

    # Фильтрация по тегу
    if selected_tag:
        query = query.filter(Post.tags.any(name=selected_tag))

    # Фильтрация по выбранным тегам
    if filter_tags:
        query = query.join(Post.tags).filter(Tag.name.in_(filter_tags))

    # Фильтрация по номеру билета
    if filter_title:
        query = query.filter_by(title=filter_title)

    # Сортировка
    if sort_by == 'title':
        query = query.order_by(Post.title.asc() if order == 'asc' else Post.title.desc())
    elif sort_by == 'level':
        query = query.order_by(Post.level.asc() if order == 'asc' else Post.level.desc())
    else:
        query = query.order_by(Post.id.asc())

    # Пагинация
    page = request.args.get('page', 1, type=int)
    posts = query.paginate(page=page, per_page=20)

    # Генерация заголовка
    if filter_title:
        dynamic_title = f"Билет: {filter_title}"
    elif selected_tag:
        dynamic_title = f"Задачи по теме: {selected_tag}"
    elif filter_tags:
        dynamic_title = f"Задачи по выбранным темам: {filter_tags}"
    else:
        dynamic_title = "Все задачи"

    all_tags = Tag.query.all()  # Все доступные теги

    return render_template(
        'posts.html',
        posts=posts,
        sort_by=sort_by,
        order=order,
        all_tags=all_tags,
        selected_tag=selected_tag,
        filter_title=filter_title,
        dynamic_title=dynamic_title
    )


@app.route('/post/<int:id>')
def show_post_detail(id):
    post = Post.query.get_or_404(id)

    def safe_b64encode(data):
        return b64encode(data).decode('utf-8') if data else None

    zad_base64 = safe_b64encode(post.zad)
    resh1_base64 = safe_b64encode(post.resh1)
    resh2_base64 = safe_b64encode(post.resh2)
    resh3_base64 = safe_b64encode(post.resh3)
    resh4_base64 = safe_b64encode(post.resh4)
    resh5_base64 = safe_b64encode(post.resh5)

    return render_template('post_detail.html', post=post, zad_base64=zad_base64, resh1_base64=resh1_base64, resh2_base64=resh2_base64, resh3_base64=resh3_base64, resh4_base64=resh4_base64,resh5_base64=resh5_base64)



@app.errorhandler(404)
def page_not_found(e):
    return render_template('404_page.html'), 404


# Страница "О проекте"
@app.route('/about')
def about():
    return render_template('about.html')


# Функция для отображения изображений в формате Base64
def to_base64(binary_data):
    return b64encode(binary_data).decode('utf-8')


# Добавление функции в шаблон
app.jinja_env.filters['to_base64'] = to_base64



# Создание папки для загрузки файлов, если она не существует
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

if __name__ == '__main__':
    app.run(debug=True, port=5000)
